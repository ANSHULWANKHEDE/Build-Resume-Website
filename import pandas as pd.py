import pandas as pd
from docx import Document

def excel_to_word(excel_file, sheet_name, word_file):
    # Read the Excel file
    df = pd.read_excel(excel_file, sheet_name=sheet_name)
    
    # Create a Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('Excel Data Sheet', 0)
    
    # Add a table with the same number of rows and columns as the DataFrame
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    
    # Add the header rows
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = col
    
    # Add the DataFrame data to the table
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iat[i, j])
    
    # Save the Word document
    doc.save(word_file)

# Example usage
excel_file = 'Automotive Wheel Aftermarket_Market sizing_July 17,2024.xlsx'  # Replace with your Excel file path
sheet_name = 'Sheet1'  # Replace with your sheet name
word_file = 'output.docx'  # Replace with your desired Word file path

excel_to_word(excel_file, sheet_name, word_file)
